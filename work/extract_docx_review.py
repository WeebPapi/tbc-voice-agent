from __future__ import annotations

import json
import sys
import zipfile
from pathlib import Path

from docx import Document
from lxml import etree


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W = "{" + W_NS + "}"
NS = {"w": W_NS}


def node_text(node) -> str:
    parts = []
    for child in node.iter():
        if child.tag in {W + "t", W + "delText"}:
            parts.append(child.text or "")
        elif child.tag == W + "tab":
            parts.append("\t")
        elif child.tag in {W + "br", W + "cr"}:
            parts.append("\n")
    return "".join(parts).strip()


def extract_paragraph(p, index: int) -> dict[str, object] | None:
    text = node_text(p)
    if not text:
        return None
    style = p.find("./w:pPr/w:pStyle", namespaces=NS)
    page_break = any(br.get(W + "type") == "page" for br in p.findall(".//w:br", namespaces=NS))
    return {
        "kind": "paragraph",
        "index": index,
        "style": style.get(W + "val") if style is not None else None,
        "page_break": page_break,
        "text": text,
    }


def extract_table(tbl, index: int) -> dict[str, object]:
    rows = []
    for tr in tbl.findall("./w:tr", namespaces=NS):
        row = []
        for tc in tr.findall("./w:tc", namespaces=NS):
            cell_parts = [node_text(p) for p in tc.findall(".//w:p", namespaces=NS)]
            row.append("\n".join(x for x in cell_parts if x))
        rows.append(row)
    return {"kind": "table", "index": index, "rows": rows}


def main() -> None:
    source = Path(sys.argv[1])
    output = Path(sys.argv[2])
    doc = Document(source)
    records = []
    for index, block in enumerate(doc.element.body.iterchildren(), 1):
        if block.tag == W + "p":
            record = extract_paragraph(block, index)
            if record:
                records.append(record)
        elif block.tag == W + "tbl":
            records.append(extract_table(block, index))

    with zipfile.ZipFile(source) as zf:
        names = set(zf.namelist())
        core = {}
        if "docProps/core.xml" in names:
            root = etree.fromstring(zf.read("docProps/core.xml"))
            core = {etree.QName(e).localname: (e.text or "") for e in root}

        comments = []
        if "word/comments.xml" in names:
            root = etree.fromstring(zf.read("word/comments.xml"))
            for comment in root.findall(".//w:comment", namespaces=NS):
                comments.append({
                    "id": comment.get(W + "id"),
                    "author": comment.get(W + "author"),
                    "date": comment.get(W + "date"),
                    "text": node_text(comment),
                })

        tracked = []
        main_root = etree.fromstring(zf.read("word/document.xml"))
        for tag_name in ("ins", "del"):
            for change in main_root.findall(f".//w:{tag_name}", namespaces=NS):
                tracked.append({
                    "type": tag_name,
                    "author": change.get(W + "author"),
                    "date": change.get(W + "date"),
                    "text": node_text(change),
                })

        auxiliary = {}
        for name in sorted(names):
            if name.startswith("word/header") and name.endswith(".xml") or name.startswith("word/footer") and name.endswith(".xml") or name in {"word/footnotes.xml", "word/endnotes.xml"}:
                root = etree.fromstring(zf.read(name))
                text = node_text(root)
                if text:
                    auxiliary[name] = text

        image_files = [name for name in names if name.startswith("word/media/")]

    result = {
        "source": str(source),
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "records": records,
        "core_properties": core,
        "comments": comments,
        "tracked_changes": tracked,
        "auxiliary_text": auxiliary,
        "image_files": image_files,
        "section_count": len(doc.sections),
    }
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps({
        "records": len(records),
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "comments": len(comments),
        "tracked_changes": len(tracked),
        "images": len(image_files),
        "sections": len(doc.sections),
    }))


if __name__ == "__main__":
    main()
